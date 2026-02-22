from __future__ import annotations

from pathlib import Path
import logging
from typing import Iterable

from .types import SourceDocument


SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}
logger = logging.getLogger(__name__)


def _safe_read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Missing pypdf dependency; cannot read PDF files.") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing python-docx dependency; cannot read DOCX files.") from exc

    doc = Document(str(path))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs).strip()


def _read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _safe_read_text(path).strip()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file type: {path}")


def _discover_files(paths: Iterable[str]) -> list[Path]:
    discovered: list[Path] = []
    for raw_path in paths:
        path = Path(raw_path).resolve()
        if not path.exists():
            continue
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES:
            discovered.append(path)
            continue
        if path.is_dir():
            for child in sorted(path.rglob("*")):
                if child.is_file() and child.suffix.lower() in SUPPORTED_SUFFIXES:
                    discovered.append(child)
    return discovered


def load_documents(paths: list[str]) -> list[SourceDocument]:
    documents: list[SourceDocument] = []
    for file_path in _discover_files(paths):
        try:
            text = _read_file(file_path)
        except Exception as exc:
            logger.warning("Source loading failed (%s)", type(exc).__name__)
            continue
        if text:
            documents.append(SourceDocument(path=str(file_path), text=text))
    return documents
